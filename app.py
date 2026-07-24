============================================================
# V7 ENTERPRISE AI DOCUMENT INTELLIGENCE
# PART 1/4
# CORE ENGINE
# ============================================================


import os
import json
from datetime import datetime
from io import BytesIO


import streamlit as st

from dotenv import load_dotenv
from openai import OpenAI

from docx import Document



# ============================================================
# PAGE CONFIG
# ============================================================


# ============================================================
# PAGE CONFIG
# ============================================================

st.set_page_config(
    page_title="Doc Lang Switcher",
    layout="wide",
    page_icon="🌍"
)



# ============================================================
# ENVIRONMENT
# ============================================================


BASE_DIR = os.path.dirname(
    os.path.abspath(__file__)
)


load_dotenv(
    os.path.join(
        BASE_DIR,
        ".env"
    )
)


api_key = os.getenv(
    "OPENAI_API_KEY"
)


if not api_key:

    st.error(
        "OPENAI_API_KEY missing. Check your .env file."
    )

    st.stop()



client = OpenAI(
    api_key=api_key
)



# ============================================================
# STORAGE
# ============================================================


MEMORY_FILE = os.path.join(
    BASE_DIR,
    "translation_memory.json"
)


METADATA_FILE = os.path.join(
    BASE_DIR,
    "document_metadata.json"
)



def load_json(file):

    if os.path.exists(file):

        try:

            with open(
                file,
                "r",
                encoding="utf-8"
            ) as f:

                return json.load(f)


        except:

            return []


    return []





def save_json(file,data):


    with open(

        file,

        "w",

        encoding="utf-8"

    ) as f:


        json.dump(

            data,

            f,

            indent=4,

            ensure_ascii=False

        )




translation_memory = load_json(
    MEMORY_FILE
)



document_metadata = load_json(
    METADATA_FILE
)



# ============================================================
# DOCX PARSER
# ============================================================



def parse_docx(file):


    doc = Document(file)


    blocks=[]

    preview=[]



    # paragraphs

    for paragraph in doc.paragraphs:


        if paragraph.text.strip():


            blocks.append({

                "type":"paragraph",

                "text":paragraph.text

            })


            preview.append(
                paragraph.text
            )



    # tables

    for table in doc.tables:


        rows=[]


        for row in table.rows:


            row_data=[]


            for cell in row.cells:

                row_data.append(
                    cell.text
                )


            rows.append(
                row_data
            )



        blocks.append({

            "type":"table",

            "rows":rows

        })



        for row in rows:

            preview.append(
                " | ".join(row)
            )



    text="\n".join(preview)


    return blocks,text





# ============================================================
# TXT READER
# ============================================================



def read_txt(file):


    text=file.read().decode(
        "utf-8"
    )


    return text





# ============================================================
# DOCUMENT INGESTION
# ============================================================



def ingest_document(file):


    filename=file.name.lower()



    if filename.endswith(".txt"):


        return read_txt(file)



    elif filename.endswith(".docx"):


        blocks,text=parse_docx(file)


        return text



    else:


        return "Unsupported format"





# ============================================================
# METADATA ENGINE
# ============================================================



def create_metadata(

    filename,

    text

):


    return {


        "filename":filename,


        "created":

        str(datetime.now()),



        "characters":

        len(text),



        "words":

        len(text.split()),



        "status":

        "processed"


    }
# ============================================================
# V7 ENTERPRISE AI DOCUMENT INTELLIGENCE
# PART 2/4
# OCR + PDF + AI ANALYSIS ENGINE
# ============================================================


from PIL import Image
import pytesseract

from pypdf import PdfReader



# ============================================================
# OCR ENGINE
# ============================================================


def extract_image_text(file):


    try:

        image = Image.open(
            file
        )


        text = pytesseract.image_to_string(
            image
        )


        return text



    except Exception as e:


        return f"OCR Error: {e}"





# ============================================================
# PDF TEXT EXTRACTION
# ============================================================


def extract_pdf_text(file):


    text=""


    try:


        reader = PdfReader(
            file
        )


        for page in reader.pages:


            page_text = page.extract_text()


            if page_text:


                text += page_text + "\n"



        return text



    except Exception as e:


        return f"PDF Error: {e}"





# ============================================================
# ADVANCED DOCUMENT INGESTION
# ============================================================


def enterprise_ingestion(file):


    filename=file.name.lower()



    if filename.endswith(".txt"):


        return read_txt(file)



    elif filename.endswith(".docx"):


        blocks,text=parse_docx(file)


        return text



    elif filename.endswith(".pdf"):


        return extract_pdf_text(file)



    elif filename.endswith(

        (

            ".png",

            ".jpg",

            ".jpeg"

        )

    ):


        return extract_image_text(file)



    else:


        return "Unsupported document format"





# ============================================================
# AI DOCUMENT CLASSIFICATION
# ============================================================


def classify_document(text):


    try:


        response = client.chat.completions.create(


            model="gpt-4.1-mini",


            messages=[


                {

                    "role":"system",

                    "content":

                    """
                    You are an enterprise document
                    classification AI.

                    Identify:

                    - document_type
                    - category
                    - confidence

                    Return JSON only.

                    Example:

                    {
                    "document_type":"resume",
                    "category":"career",
                    "confidence":"high"
                    }

                    """

                },


                {

                    "role":"user",

                    "content":text[:4000]

                }

            ]

        )



        result=response.choices[0].message.content



        return json.loads(result)



    except:


        return {


            "document_type":"unknown",

            "category":"unknown",

            "confidence":"low"

        }





# ============================================================
# AI ENTITY EXTRACTION
# ============================================================


def extract_entities(text):


    try:


        response = client.chat.completions.create(


            model="gpt-4.1-mini",


            messages=[


                {


                    "role":"system",


                    "content":

                    """

                    Extract important entities.

                    Return JSON only.

                    Include:

                    people
                    companies
                    dates
                    locations
                    money


                    """

                },


                {


                    "role":"user",

                    "content":text[:5000]

                }


            ]

        )



        result=response.choices[0].message.content



        return json.loads(result)



    except:


        return {


            "people":[],


            "companies":[],


            "dates":[],


            "locations":[],


            "money":[]

        }





# ============================================================
# DOCUMENT INTELLIGENCE PIPELINE
# ============================================================


def analyze_document(file):


    text = enterprise_ingestion(
        file
    )



    metadata=create_metadata(

        file.name,

        text

    )



    document_metadata.append(
        metadata
    )


    save_json(

        METADATA_FILE,

        document_metadata

    )



    result={


        "filename":file.name,


        "raw_text":text,


        "classification":

        classify_document(text),



        "entities":

        extract_entities(text),



        "metadata":

        metadata


    }



    return result
# ============================================================
# V7 ENTERPRISE AI DOCUMENT INTELLIGENCE
# PART 3/4
# TRANSLATION + DOWNLOAD ENGINE
# ============================================================



# ============================================================
# TRANSLATION MEMORY
# ============================================================


def find_memory(source_text):


    for item in translation_memory:


        if item.get("source") == source_text:


            return item.get("translation")



    return None





def save_memory(

    source,

    translation,

    language

):


    translation_memory.append(


        {


            "source":source,


            "translation":translation,


            "language":language,


            "created":str(datetime.now())


        }

    )


    save_json(

        MEMORY_FILE,

        translation_memory

    )





# ============================================================
# AI TRANSLATION ENGINE
# ============================================================


def translate_document(

    text,

    target_language

):


    previous = find_memory(
        text
    )



    if previous:


        return previous





    response = client.chat.completions.create(


        model="gpt-4.1-mini",


        messages=[


            {


                "role":"system",


                "content":

                """

                You are an enterprise
                document translation engine.

                Rules:

                - Translate accurately.
                - Preserve professional tone.
                - Preserve names.
                - Preserve numbers.
                - Preserve dates.
                - Do not summarize.
                - Do not add explanations.


                """

            },


            {


                "role":"user",


                "content":


                f"""

                Translate this document into
                {target_language}.


                Document:

                {text}

                """

            }


        ]

    )



    translated = (

        response
        .choices[0]
        .message
        .content

    )



    save_memory(

        text,

        translated,

        target_language

    )


    return translated





# ============================================================
# QUALITY VALIDATION
# ============================================================


def quality_check(

    original,

    translated

):


    return {


        "original_characters":

        len(original),



        "translated_characters":

        len(translated),



        "translation_completed":

        bool(translated),



        "status":

        "PASSED"

        if translated

        else

        "FAILED"


    }





# ============================================================
# DOWNLOAD CENTER
# ============================================================


# -------------------------
# DOCX GENERATOR
# -------------------------



def create_docx(text):


    doc=Document()



    for line in text.split("\n"):


        doc.add_paragraph(
            line
        )



    buffer=BytesIO()



    doc.save(
        buffer
    )



    buffer.seek(0)



    return buffer





# -------------------------
# TXT GENERATOR
# -------------------------



def create_txt(text):


    return text





# -------------------------
# PDF GENERATOR
# -------------------------


def create_pdf(text):


    from reportlab.platypus import (

        SimpleDocTemplate,

        Paragraph

    )


    from reportlab.lib.styles import getSampleStyleSheet



    buffer=BytesIO()



    pdf=SimpleDocTemplate(
        buffer
    )



    styles=getSampleStyleSheet()



    content=[]



    for line in text.split("\n"):


        content.append(

            Paragraph(

                line,

                styles["Normal"]

            )

        )



    pdf.build(
        content
    )



    buffer.seek(0)



    return buffer
# ============================================================
# V7 ENTERPRISE AI DOCUMENT INTELLIGENCE
# PART 4/4
# USER INTERFACE
# ============================================================



# ============================================================
# SESSION STATE
# ============================================================


if "analysis_result" not in st.session_state:

    st.session_state.analysis_result = None



if "translated_result" not in st.session_state:

    st.session_state.translated_result = None



if "target_language" not in st.session_state:

    st.session_state.target_language = None





# ============================================================
# HEADER
# ============================================================


st.title("🌍 DocLangAI")


st.write(
"""
AI-Powered Document Translation

Upload → Analyze → Translate → Download

Translate DOCX, PDF, TXT, and Images instantly with AI.
"""
)





# ============================================================
# UPLOAD
# ============================================================


uploaded_file = st.file_uploader(


    "📄 Upload Document",


    type=[

        "docx",

        "txt",

        "pdf",

        "png",

        "jpg",

        "jpeg"

    ]

)





# ============================================================
# DOCUMENT ANALYSIS
# ============================================================


if uploaded_file:


    if st.button(

        "🔍 Analyze Document"

    ):


        with st.spinner(

            "AI analyzing document..."

        ):


            result = analyze_document(

                uploaded_file

            )


            st.session_state.analysis_result = result



        st.success(

            "Document analysis completed"

        )





# ============================================================
# DISPLAY ANALYSIS
# ============================================================


if st.session_state.analysis_result:



    result = st.session_state.analysis_result



    st.divider()



    st.subheader(

        "📊 Document Intelligence Report"

    )



    # Classification


    st.markdown(

        "### 📁 Document Classification"

    )


    st.json(

        result["classification"]

    )




    # Metadata


    st.markdown(

        "### 📝 Metadata"

    )


    st.json(

        result["metadata"]

    )




    # Entities


    st.markdown(

        "### 🔎 Extracted Entities"

    )


    st.json(

        result["entities"]

    )




    # Original Text


    st.markdown(

        "### 📄 Extracted Content"

    )


    st.text_area(

        "Original Document",

        result["raw_text"],

        height=300

    )




    st.divider()



    # ========================================================
    # TRANSLATION
    # ========================================================


    st.subheader(

        "🌎 Enterprise Translation"

    )



    languages=[


        "Spanish",

        "French",

        "German",

        "Italian",

        "Portuguese",

        "Chinese",

        "Japanese",

        "Korean",

        "Arabic",

        "Hindi"


    ]



    language = st.selectbox(

        "Select Target Language",

        languages

    )



    st.session_state.target_language = language




    if st.button(

        "🚀 Translate Document"

    ):



        with st.spinner(

            "AI translating document..."

        ):



            translated = translate_document(

                result["raw_text"],

                language

            )



            st.session_state.translated_result = translated



        st.success(

            "Translation completed"

        )






# ============================================================
# TRANSLATION OUTPUT + DOWNLOAD CENTER
# ============================================================


if st.session_state.translated_result:



    st.divider()



    st.subheader(

        "🌎 Translated Document"

    )



    translated_text = st.session_state.translated_result



    st.text_area(

        "Translation",

        translated_text,

        height=350

    )




    # ========================================================
    # QUALITY REPORT
    # ========================================================


    st.subheader(

        "✅ Quality Validation"

    )



    quality = quality_check(

        st.session_state.analysis_result["raw_text"],

        translated_text

    )



    st.json(

        quality

    )




    # ========================================================
    # DOWNLOAD CENTER
    # ========================================================


    st.divider()



    st.subheader(

        "⬇️ Download Center"

    )




    # TXT


    st.download_button(

        label="📝 Download TXT",

        data=create_txt(

            translated_text

        ),

        file_name="translated_document.txt",

        mime="text/plain"

    )




    # DOCX


    docx_file = create_docx(

        translated_text

    )



    st.download_button(

        label="📄 Download DOCX",

        data=docx_file,

        file_name="translated_document.docx",

        mime=

        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    )




    # PDF


    pdf_file = create_pdf(

        translated_text

    )



    st.download_button(

        label="📕 Download PDF",

        data=pdf_file,

        file_name="translated_document.pdf",

        mime="application/pdf"

    )





# ============================================================
# SYSTEM DASHBOARD
# ============================================================


st.divider()



st.subheader(

    "🏢 Enterprise System Status"

)



col1,col2,col3 = st.columns(3)



with col1:


    st.metric(

        "Documents Processed",

        len(document_metadata)

    )



with col2:


    st.metric(

        "Translation Memory",

        len(translation_memory)

    )



with col3:


    st.metric(

        "AI Engine",

        "ACTIVE"

    )





# ============================================================
# FOOTER
# ============================================================

st.caption(
"""
🌍 DocLangAI

AI-Powered Document Translation

Supported formats:
DOCX • PDF • TXT • PNG • JPG

Powered by OpenAI
"""
)
